from django.shortcuts import render
from django.http import HttpResponse, HttpResponseRedirect, FileResponse
from django.urls import reverse, reverse_lazy
from django.template import loader
from django.db import connection
from django.contrib.auth.forms import UserCreationForm
from django.views import generic

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
import webbrowser
import platform
import os
import smtplib
import ssl
from email.message import EmailMessage
from datetime import date, timedelta, datetime
from decimal import Decimal
import math
import docx

from .OrderFunctions import OrderFunctions as OrderFunc
from .OrdersDD import OrdersDD

from .models import Batchcost, Materialcost, Materialinventory, Materialtesting, Ordersheetmachine1, Ordersheetmachine2, Ordersheetmachine3, Productinventory, Productprofiles, Colour, Employees, Profileaverages, picsum, picsumForm, cisEmailSubject

numberOfMachines = 3
num_of_materials = 11

###########################################################HOME PAGE
def index(request):
    return render(request, 'CWBDataApp/index.html')

###########################################################BATCH COST TRACKING
def BatchCostTracking(request):
    allColours = Colour.objects.all()
    allMaterials= Materialinventory.objects.all()
    allProfiles= Productprofiles.objects.all()

    if request.method == 'POST':

        form = request.POST

        try:
            #Check if batch already exists, if so stop and send an error message
            if Batchcost.objects.get(pk=request.POST['newBatch']):
                return render(request, 'CWBDataApp/BatchCostTracking.html', {'batch':form, 'allMaterials':allMaterials, 'allColours':allColours, 'allProfiles':allProfiles, 'range':range(1, num_of_materials), 'error_message' : "Batch already exists, please enter a new batch. If you wish to update a batch talk to an admin",})
        except:

            #Calculate total price, weight and price/pound
            cost = 0
            weight = 0
            for i in range(1, num_of_materials):
                cost += int(form['weight'+ str(i)]) * Decimal(form['value'+str(i)])
            cost += (Decimal(form['colourWeight']) * Decimal(form['colourPrice'])) + (Decimal(form['foamWeight']) * Decimal(form['foamPrice']))
            cost=round(cost, 2)

            for i in range(1, num_of_materials):
                weight += int(form['weight'+str(i)])
            weight += Decimal(form['colourWeight']) + Decimal(form['foamWeight'])
            #Check if weight was entered as 0, throw error if it is
            if weight == 0:
                return render(request, 'CWBDataApp/BatchCostTracking.html', {'batch':form, 'allMaterials':allMaterials, 'allColours':allColours, 'allProfiles':allProfiles, 'range':range(1, num_of_materials), 'error_message':"Total weight cannot be zero, please add a value to weight1",})

            price = round(cost/weight, 2)
            #Check if theres enough boxes before we subtract from inventory
            for i in range(1, num_of_materials):
                try:
                    if form['material'+str(i)] != 'None':
                        box = Materialinventory.objects.get(pk=form['material'+str(i)])

                        if box.numberofboxes < Decimal(form['numofBoxes'+str(i)]):
                            return render(request, 'CWBDataApp/BatchCostTracking.html', {'batch':form, 'allMaterials':allMaterials, 'allColours':allColours, 'allProfiles':allProfiles, 'range':range(1, num_of_materials), 'error_message' : form["material"+str(i)]+" does no have enough boxes for Material"+str(i)+"",})
                except:
                    return render(request, 'CWBDataApp/BatchCostTracking.html', {'batch':form, 'allMaterials':allMaterials, 'allColours':allColours, 'allProfiles':allProfiles, 'range':range(1, num_of_materials), 'error_message' : "Couldn't find box in inventory",})

            materials = {}
            numberofboxes = 0
            #Remove boxes from inventory if material value given was from INVENTORY
            #If a custom material was entered do not remove from inventory
            for i in range(1, num_of_materials):
                #Sum total number of boxes
                try:
                    if form['material'+str(i)] != 'None':
                        box = Materialinventory.objects.get(pk=form['material'+str(i)])
                        box.numberofboxes -= Decimal(form['numofBoxes'+str(i)])
                        box.save()
                        if box.numberofboxes <=0:
                            box.delete()
                        materials['material'+str(i)] = form['material'+str(i)]
                        numberofboxes += Decimal(form['numofBoxes'+str(i)])
                    elif form['material'+str(i)] == 'None' and form['material'+str(i)+'2'] != '':
                        materials['material'+str(i)] = form['material'+str(i)+'2']
                        numberofboxes += Decimal(form['numofBoxes'+str(i)])
                    else:
                        materials['material'+str(i)] = 'None'
                except:
                    return render(request, 'CWBDataApp/BatchCostTracking.html', {'batch':form, 'allMaterials':allMaterials, 'allColours':allColours, 'allProfiles':allProfiles, 'range':range(1, num_of_materials), 'dateToday':str(date.today), 'error_message' : form["material"+str(i)]+" does not have enough boxes for Material"+str(i)+"",})

            sup_object = Materialcost.objects.get(pk='Free')

            #If batch exists in inventory then add the number of boxes to it, this SHOULD NOT HAPPEN THOUGH
            if Materialinventory.objects.filter(materialname=form['newBatch']).exists():
                materialInv = Materialinventory.objects.get(pk=form['newBatch'])
                materialInv.numberofboxes += numberofboxes
            else:
                materialInv = Materialinventory(materialname=form['newBatch'].upper(),
                                                supplier=sup_object,
                                                numberofboxes=numberofboxes,
                                                locations='None',
                                                premixed='True',
                                                priceperpound=price)
                materialInv.save()

            #Create new batch entry if this isnt a premix
            if form['premix'] == 'No':
                batch = Batchcost(batchname=form['newBatch'].upper(),
                                        batchdate=form['batchDate'],
                                        totalcost=cost,
                                        totalweight=weight,
                                        priceperpound=price,
                                        material1=materials['material1'],
                                        weight1=form['weight1'],
                                        value1=form['value1'],
                                        material2=materials['material2'],
                                        weight2=form['weight2'],
                                        value2=form['value2'],
                                        material3=materials['material3'],
                                        weight3=form['weight3'],
                                        value3=form['value3'],
                                        material4=materials['material4'],
                                        weight4=form['weight4'],
                                        value4=form['value4'],
                                        material5=materials['material5'],
                                        weight5=form['weight5'],
                                        value5=form['value5'],
                                        material6=materials['material6'],
                                        weight6=form['weight6'],
                                        value6=form['value6'],
                                        material7=materials['material7'],
                                        weight7=form['weight7'],
                                        value7=form['value7'],
                                        material8=materials['material8'],
                                        weight8=form['weight8'],
                                        value8=form['value8'],
                                        material9=materials['material9'],
                                        weight9=form['weight9'],
                                        value9=form['value9'],
                                        material10=materials['material10'],
                                        weight10=form['weight10'],
                                        value10=form['value10'],
                                        colour=form['colour'],
                                        colourweight=form['colourWeight'],
                                        colourvalue=form['colourPrice'],
                                        foamweight=form['foamWeight'],
                                        foamvalue=form['foamPrice'],
                                        totalshredweight=form['shredWeight'],
                                        profile=form['profile']
                                        )
                batch.save()
            return render(request, 'CWBDataApp/BatchCostTracking.html', {'allMaterials':allMaterials, 'allColours':allColours, 'allProfiles':allProfiles, 'range':range(1, num_of_materials), 'dateToday':str(date.today), 'dataAcceptedMessage':"Batch Successfully Submitted and material used was removed from inventory"})


    return render(request, 'CWBDataApp/BatchCostTracking.html', {'allMaterials':allMaterials, 'allColours':allColours, 'allProfiles':allProfiles, 'range':range(1, num_of_materials), 'dateToday':date.today().isoformat()})

###########################################################UPDATE BATCH
def UpdateBatch(request):

    allBatchs = Batchcost.objects.all()

    if request.method == 'POST':
        form = request.POST
        batchname = ''

        if form['batch'] == 'None':
            batchname = form['batch2']
        else:
            batchname = form['batch']

        try:
            #Try and grab existing batch, throws eception if failed
            batch = Batchcost.objects.get(pk=batchname)
            batchDate = batch.batchdate.isoformat()
            #Get values need for batch
            allProfiles = Productprofiles.objects.all()
            allColours = Colour.objects.all()
            allMaterials= Materialinventory.objects.all()
            return render(request, 'CWBDataApp/UpdateBatch.html', {'allBatchs':allBatchs, 'allProfiles':allProfiles, 'allColours':allColours, 'allMaterials':allMaterials, 'range':range(1, num_of_materials), 'batch':batch, 'batchDate':batchDate})
        except:
            return render(request, 'CWBDataApp/UpdateBatch.html', {'allBatchs':allBatchs})

    return render(request, 'CWBDataApp/UpdateBatch.html', {'allBatchs':allBatchs})

def ChangeBatch(request):
    allBatchs = Batchcost.objects.all()

    if request.method == 'POST':

        form = request.POST

        allProfiles = Productprofiles.objects.all()
        allColours = Colour.objects.all()
        allMaterials= Materialinventory.objects.all()

        batch = Batchcost.objects.get(pk=form['newBatch'].upper())

        #Calculate total price, weight and price/pound
        cost = 0
        weight = 0
        weightDict = {}
        valueDict = {}
        for i in range(1, num_of_materials):

            #Check if empty values exists, fill in blanks with 0
            if form['weight'+ str(i)] == '':
                weightDict['weight'+ str(i)] = 0
            else:
                weightDict['weight'+ str(i)] = form['weight'+ str(i)]

            if form['value'+ str(i)] == '':
                valueDict['value'+ str(i)] = 0
            else:
                valueDict['value'+ str(i)] = form['value'+ str(i)]

            #Sum up cost of materials
            cost += int(weightDict['weight'+ str(i)]) * Decimal(valueDict['value'+ str(i)])

        #Add cost of colour and foam
        cost += (Decimal(form['colourWeight']) * Decimal(form['colourPrice'])) + (Decimal(form['foamWeight']) * Decimal(form['foamPrice']))
        cost=round(cost, 2)

        for i in range(1, num_of_materials):
            weight += int(weightDict['weight'+str(i)])
        weight += Decimal(form['colourWeight']) + Decimal(form['foamWeight'])
        #Check if weight was entered as 0, throw error if it is
        if weight == 0:
            return render(request, 'CWBDataApp/UpdateBatch.html', {'allBatchs':allBatchs, 'allProfiles':allProfiles, 'allColours':allColours, 'allMaterials':allMaterials, 'range':range(1, num_of_materials), 'batch':batch, 'batchDate':form['batchDate'], 'error_message':'Weights cannot be zero, please enter a new value' })

        price = round(cost/weight, 2)
        #Check if theres enough boxes before we subtract from inventory
        for i in range(1, num_of_materials):
            try:
                if form['material'+str(i)] != 'None':
                    box = Materialinventory.objects.get(pk=form['material'+str(i)])

                    if box.numberofboxes < Decimal(form['numofBoxes'+str(i)]):
                        return render(request, 'CWBDataApp/UpdateBatch.html', {'allBatchs':allBatchs, 'allMaterials':allMaterials, 'allColours':allColours, 'allProfiles':allProfiles, 'range':range(1, num_of_materials), 'batch':batch, 'batchDate':form['batchDate'], 'error_message' : form["material"+str(i)]+" does no have enough boxes for Material"+str(i)+"",})
            except:
                return render(request, 'CWBDataApp/BatchCostTracking.html', {'allBatchs':allBatchs, 'allMaterials':allMaterials, 'allColours':allColours, 'allProfiles':allProfiles, 'range':range(1, num_of_materials), 'batch':batch, 'batchDate':form['batchDate'], 'error_message' : "Couldn't find box in inventory",})

        materials = {}
        numberofboxes = 0
        #Remove boxes from inventory if material value given was from INVENTORY
        #If a custom material was entered do not remove from inventory
        for i in range(1, num_of_materials):
            #Sum total number of boxes
            try:
                #Grab current materials weight from inventory
                weightBatch = getattr(batch, 'weight'+str(i))

                #Check if a new material was added or if the weight increased from the current weight, this indicates we should remove boxes from inventory
                if form['material'+str(i)] != 'None' or int(form['weight'+str(i)]) != weightBatch :
                    if  Materialinventory.objects.filter(pk=form['material'+str(i)]).exists():
                        box = Materialinventory.objects.get(pk=form['material'+str(i)])
                        box.numberofboxes -= Decimal(form['numofBoxes'+str(i)])
                        box.save()
                        if box.numberofboxes <=0:
                            box.delete()
                    materials['material'+str(i)] = form['material'+str(i)]
                    numberofboxes += Decimal(form['numofBoxes'+str(i)])
                elif form['material'+str(i)+'2'] != '':
                    materials['material'+str(i)] = form['material'+str(i)+'2']
                    numberofboxes += Decimal(form['numofBoxes'+str(i)])
                else:
                    materials['material'+str(i)] = 'None'
            except:
                return render(request, 'CWBDataApp/UpdateBatch.html', {'allBatchs':allBatchs,'allMaterials':allMaterials, 'allColours':allColours, 'allProfiles':allProfiles, 'range':range(1, num_of_materials), 'batch':batch, 'batchDate':form['batchDate'], 'error_message' : form["material"+str(i)]+" does not have enough boxes for Material"+str(i)+"",})

        sup_object = Materialcost.objects.get(pk='Free')

        #If batch exists in inventory then add the number of boxes to it, this SHOULD HAPPEN
        if Materialinventory.objects.filter(materialname=form['newBatch']).exists():
            materialInv= Materialinventory.objects.get(pk=form['newBatch'])
            materialInv.numberofboxes = numberofboxes
            materialInv.priceperpound = price
            materialInv.save()
        else:
            materialInv = Materialinventory(materialname=form['newBatch'].upper(),
                                            supplier=sup_object,
                                            numberofboxes=numberofboxes,
                                            locations='None',
                                            premixed='True',
                                            priceperpound=price)
            materialInv.save()

        #Create new batch entry if this isnt a premix
        if form['premix'] == 'No':

            batch.batchdate=form['batchDate']
            batch.totalcost=cost
            batch.totalweight=weight
            batch.priceperpound=price
            batch.material1=materials['material1']
            batch.weight1=weightDict['weight1']
            batch.value1=valueDict['value1']
            batch.material2=materials['material2']
            batch.weight2=weightDict['weight2']
            batch.value2=valueDict['value2']
            batch.material3=materials['material3']
            batch.weight3=weightDict['weight3']
            batch.value3=valueDict['value3']
            batch.material4=materials['material4']
            batch.weight4=weightDict['weight4']
            batch.value4=valueDict['value4']
            batch.material5=materials['material5']
            batch.weight5=weightDict['weight5']
            batch.value5=valueDict['value5']
            batch.material6=materials['material6']
            batch.weight6=weightDict['weight6']
            batch.value6=valueDict['value6']
            batch.material7=materials['material7']
            batch.weight7=weightDict['weight7']
            batch.value7=valueDict['value7']
            batch.material8=materials['material8']
            batch.weight8=weightDict['weight8']
            batch.value8=valueDict['value8']
            batch.material9=materials['material9']
            batch.weight9=weightDict['weight9']
            batch.value9=valueDict['value9']
            batch.material10=materials['material10']
            batch.weight10=weightDict['weight10']
            batch.value10=valueDict['value10']
            batch.colour=form['colour']
            batch.colourweight=form['colourWeight']
            batch.colourvalue=form['colourPrice']
            batch.foamweight=form['foamWeight']
            batch.foamvalue=form['foamPrice']
            batch.totalshredweight=form['shredWeight']
            batch.profile=form['profile']

            batch.save()
        return render(request, 'CWBDataApp/UpdateBatch.html', {'allBatchs':allBatchs, 'dataAcceptedMessage':"Batch "+form['newBatch']+" was successfully updated"})



    return render(request, 'CWBDataApp/UpdateBatch.html', {'allBatchs':allBatchs})
###########################################################BATCH COST TRACKING QUERY
def BatchCostQuery(request):

    if request.method == 'POST':
        form = request.POST

        try:
            #Grab batch if it exists, throw error if it doesn't
            batch = Batchcost.objects.get(pk=form['searchBatch'].upper())
            dict= {}
            for i in range(1, num_of_materials):
                dict['material'+str(i)] = getattr(batch, 'material'+str(i))
                dict['weight'+str(i)] = getattr(batch, 'weight'+str(i))
                dict['value'+str(i)] = getattr(batch, 'value'+str(i))
            return render(request, 'CWBDataApp/BatchCostQuery.html', {'batch' : batch, 'range':range(1,num_of_materials), 'dict':dict.items(),})
        except:
            return render(request, 'CWBDataApp/BatchCostQuery.html', {'error_message' : "Batch does not exist, please enter a valid batch",})
    return render(request, 'CWBDataApp/BatchCostQuery.html')

###########################################################Batch Cost Excel File Download
def BatchCostExcel(request):
    query = str(Batchcost.objects.all().query)
    df = pd.read_sql_query(query, connection)
    today = str(date.today())

    df.to_excel(r'./CWBDataApp/BatchCostTracking.xlsx', index=False)

    with open(r'./CWBDataApp/BatchCostTracking.xlsx', 'rb') as fh:
        response = HttpResponse(fh.read(), content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        response['Content-Disposition'] = 'attachment; filename=BatchCostTracking-'+today+'.xlsx'
    os.remove(r'./CWBDataApp/BatchCostTracking.xlsx')
    return response

    return render(request, 'CWBDataApp/BatchCostQuery.html')

###########################################################MATERIAL TESTING
def MaterialTesting(request):
    allProfiles= Productprofiles.objects.all()
    allEmployees= Employees.objects.all()

    if request.method == 'POST':

        form = request.POST

        try:
            Batchcosttracking.objects.get(pk=form['testName'].upper())
        except:
            return render(request, 'CWBDataApp/MaterialTesting.html', {'allProfiles':allProfiles, 'allEmployees':allEmployees, 'error_message' : "Batch doesn't exist, please enter a batch first before entering a test.",})

        try:
            testName = Batchcosttracking.objects.get(pk=form['testName'].upper())
            Materialtesting.objects.get(testname=testName, testnumber=int(form['testNum']))
            return render(request, 'CWBDataApp/MaterialTesting.html', {'allProfiles':allProfiles, 'allEmployees':allEmployees, 'error_message' : "Test number is already used for this batch, please choose a different test number",})
        except:
            if int(form['machinetimeUsed']) == 1:
                costofmaterial=150
                costoflabour=100
            elif int(form['machinetimeUsed']) == 2:
                costofmaterial=200
                costoflabour=200
            elif int(form['machinetimeUsed']) == 3:
                costofmaterial=400
                costoflabour=300
            elif int(form['machinetimeUsed']) == 4:
                costofmaterial=450
                costoflabour=300
            elif int(form['machinetimeUsed']) == 5:
                costofmaterial=500
                costoflabour=300
            elif int(form['machinetimeUsed']) == 6:
                costofmaterial=600
                costoflabour=350
            elif int(form['machinetimeUsed']) == 8:
                costofmaterial=900
                costoflabour=400
            elif Decimal(form['machinetimeUsed']) == 0.5:
                costofmaterial=150
                costoflabour=50
            else:
                costofmaterial=1200
                costoflabour=600

            totalcost = costoflabour + costofmaterial
            testName = Batchcosttracking.objects.get(pk=form['testName'].upper())

            test=Materialtesting(projectnumber=form['projNum'],
                                 testname=testName,
                                 testnumber=form['testNum'],
                                 testdate=form['testDate'],
                                 labourused=form['labourUsed'],
                                 machinetimeused=form['machinetimeUsed'],
                                 productionline=form['productionLine'],
                                 materialstested=form['materialsTested'],
                                 othermaterialstested=form['othermaterialsTested'],
                                 moulds=form['moulds'],
                                 reasonfortest=form['reasonForTest'],
                                 expectedresults=form['expectedResults'],
                                 difficultiesencountered=form['difficultiesencountered'],
                                 nextstep=form['nextStep'],
                                 estimatedcostofmaterial=costofmaterial,
                                 estimatedcostoflabour=costoflabour,
                                 totalcostoftest=totalcost
                                 )
            test.save()
            return render(request, 'CWBDataApp/MaterialTesting.html', {'allProfiles':allProfiles, 'allEmployees':allEmployees, 'dataAcceptedMessage':"Test Successfully Submitted"})
    return render(request, 'CWBDataApp/MaterialTesting.html', {'allProfiles':allProfiles, 'allEmployees':allEmployees})

###########################################################MATERIAL TESTING Populate
def MaterialTestingPopulate(request):
    allProfiles= Productprofiles.objects.all()

    if request.method == 'POST':

        form = request.POST

        if int(form['testNum']) > 1:
            try:
                Batchcosttracking.objects.get(pk=form['testName'].upper())
            except:
                return render(request, 'CWBDataApp/MaterialTesting.html', {'allProfiles':allProfiles, 'error_message' : "Batch doesn't exist, please enter a batch first before entering a test.",})

            try:
                testName = Batchcosttracking.objects.get(pk=form['testName'].upper())
                Materialtesting.objects.get(testname=testName, testnumber=int(form['testNum']))
                return render(request, 'CWBDataApp/MaterialTesting.html', {'allProfiles':allProfiles, 'error_message' : "Test number is already used for this batch, please choose a different test number",})
            except:

                testName = Batchcosttracking.objects.get(pk=form['testName'].upper())
                prev_test = Materialtesting.objects.get(testname=testName, testnumber=int(form['testNum'])-1)

                pop_data = {}
                entered_data = {}
                for key, value in form.items():
                    if value == '':
                        pop_data[key.lower()] = getattr(prev_test, key.lower())
                    else:
                        entered_data[key.lower()] = value
                return render(request, 'CWBDataApp/MaterialTesting.html', {'allProfiles':allProfiles, 'dataAcceptedMessage':"Data Populated", 'pop_data':pop_data, 'entered_data':entered_data})

                return render(request, 'CWBDataApp/MaterialTesting.html', {'allProfiles':allProfiles, 'error_message' : "Something went wrong, can't populate data",})
        else:
            return render(request, 'CWBDataApp/MaterialTesting.html', {'allProfiles':allProfiles, 'error_message':"Cannot populate data on the first test"})
    return render(request, 'CWBDataApp/MaterialTesting.html', {'allProfiles':allProfiles})

###########################################################MATERIAL TESTING QUERY
def MaterialTestQuery(request):

    if request.method == 'POST':

        form = request.POST

        try:
            test = Batchcosttracking.objects.get(pk=form['testName'].upper())
            testName = test.batchname
            allTests = Materialtesting.objects.filter(testname=test)
            if not allTests:
                return render(request, 'CWBDataApp/MaterialTestQuery.html', {'error_message':"No test exists for this batch"})
            return render(request, 'CWBDataApp/MaterialTestQuery.html', {'allTests':allTests, 'testName':testName})
        except:
            return render(request, 'CWBDataApp/MaterialTestQuery.html', {'error_message':"Batch does not exist, please enter the batch in batch cost tracking"})
    return render(request, 'CWBDataApp/MaterialTestQuery.html')

###########################################################Material Test Excel File Download
def MaterialTestExcel(request):
    query = str(Materialtesting.objects.all().query)
    df = pd.read_sql_query(query, connection)
    today = str(date.today())

    df.to_excel(r'./CWBDataApp/MaterialTesting.xlsx', index=False)

    with open(r'./CWBDataApp/MaterialTesting.xlsx', 'rb') as fh:
        response = HttpResponse(fh.read(), content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        response['Content-Disposition'] = 'attachment; filename=MaterialTesting-'+today+'.xlsx'
    os.remove(r'./CWBDataApp/MaterialTesting.xlsx')
    return response

    return render(request, 'CWBDataApp/MaterialTestQuery.html')


###########################################################PRODUCT INVENTORY
def ProductInventory(request):

    allProfiles= Productprofiles.objects.all()
    allColours = Colour.objects.all()

    if request.method == 'POST':
        form = request.POST
        if Productinventory.objects.filter(productname=form['prodName'], colour=form['colour']).exists():

            return render(request, 'CWBDataApp/ProductInventory.html', {'allProfiles':allProfiles, 'allColours':allColours, 'error_message' : "Product already exists in database, please enter a new product. If you wish to update a product scroll down",})
        elif form['numSkids'] == '' or Decimal(form['numSkids']) <= 0:
            return render(request, 'CWBDataApp/ProductInventory.html', {'allProfiles':allProfiles, 'allColours':allColours, 'error_message' : "Cannot enter a product with zero skids",})
        else:
            #Check to see if there is an order that contains this product
            Order = OrderFunc()
            orderDict = Order.FindOrder(form['prodName'], form['colour'])
            prodProf = Productprofiles.objects.get(pk=form['prodName'])
            #If an order exists containing this product, then update its inventorized pieces
            if orderDict:
                prodAv = Profileaverages.objects.get(pk=form['prodName'])
                numPieces = int(Decimal(form['numSkids']) * Decimal(prodProf.pcsperskid)) + int(form['pcs'])
                Order.UpdateOrderInventory(orderDict['order'], numPieces, orderDict['orderSheet'], prodProf, prodAv)

            #Turn numofskids and extra peices into a decimal number
            numSkids = Decimal(form['numSkids']) + (int(form['pcs']) / Decimal(prodProf.pcsperskid))

            newProd = Productinventory(productname=form['prodName'],
                                       colour=form['colour'],
                                       embossed=form['embossed'],
                                       doublesided=form['doubleSide'],
                                       numberofskids=numSkids)
            newProd.save()
            return render(request, 'CWBDataApp/ProductInventory.html', {'allProfiles':allProfiles, 'allColours':allColours, 'dataAcceptedMessage':"Data Successfully Submitted"})
    return render(request, 'CWBDataApp/ProductInventory.html', {'allProfiles':allProfiles, 'allColours':allColours})

###########################################################PRODUCT INVENTORY UPDATE
def ProductInventoryUpdate(request):
    allProduct = Productinventory.objects.all()
    allColours = Colour.objects.all()

    if request.method == 'POST':
        form = request.POST

        if Productinventory.objects.filter(pk=form['prodName']).exists():
            prod_object = Productinventory.objects.get(pk=form['prodName'])
            return render(request, 'CWBDataApp/ProductInventoryUpdate.html', {'allProduct':allProduct, 'allColours':allColours, 'Forms':True, 'prod_object':prod_object})
        else:
            return render(request, 'CWBDataApp/ProductInventoryUpdate.html', {'allProduct':allProduct, 'allColours':allColours, 'error_message': ""+ form['colour'] + " " + form['prodName'] +" does not exist in inventory. If you wish to enter its data, press the Add Product tab"})
    return render(request, 'CWBDataApp/ProductInventoryUpdate.html', {'allProduct':allProduct})

###########################################################PRODUCT INVENTORY UPDATE SKIDS
def ProductInventoryUpdateSkids(request):
    allProduct = Productinventory.objects.all()

    if request.method == 'POST':
        form = request.POST

        #if our colour has changed then update the inventory and ordersheets
        if form['newcolour'] != form['colour']:
            #Check if a product already exists with this new colour
            if Productinventory.objects.filter(productname=form['prodName'], colour=form['newcolour']).exists():
                return render(request, 'CWBDataApp/ProductInventoryUpdate.html', {'allProduct':allProduct, 'error_message': 'A Product with this profile and colour already exists'})

            prod = Productinventory.objects.get(productname=form['prodName'], colour=form['colour'])
            prod.colour = form['newcolour']
            prod.save()

            prodProf = Productprofiles.objects.get(pk=form['prodName'])
            prodAv = Profileaverages.objects.get(pk=form['prodName'])

            pieces = prod.numberofskids * Decimal(prodProf.pcsperskid)

            #See if an order exists with the product old colour and substract that from inventory
            Order = OrderFunc()
            orderDict = Order.FindOrder(form['prodName'], form['colour'])
            #If an order exists then update its inventorized pieces
            if orderDict:
                Order.UpdateOrderInventory(orderDict.get('order'), (-1 * pieces), orderDict.get('orderSheet'), prodProf, prodAv)

            #See if an order exists with the producs new colour and add that to inventory
            Order = OrderFunc()
            orderDict = Order.FindOrder(form['prodName'], form['newcolour'])
            #If an order exists then update its inventorized pieces
            if orderDict:
                Order.UpdateOrderInventory(orderDict.get('order'), pieces, orderDict.get('orderSheet'), prodProf, prodAv)

        #Check if any new skids were added, if numskids is 0 then skip
        if form['numSkids'] != '' and form['numSkids'] != 0:

            if Productinventory.objects.filter(productname=form['prodName'], colour=form['newcolour']).exists():
                prod = Productinventory.objects.get(productname=form['prodName'], colour=form['newcolour'])

                numSkids = Decimal(form['numSkids'])

                prodProf = Productprofiles.objects.get(pk=form['prodName'])
                prodAv = Profileaverages.objects.get(pk=form['prodName'])
                differenceInPieces = numSkids * Decimal(prodProf.pcsperskid)
                #See if an order exists with this product
                Order = OrderFunc()
                orderDict = Order.FindOrder(form['prodName'], form['newcolour'])
                #If an order exists then update its inventorized pieces
                if orderDict:
                    Order.UpdateOrderInventory(orderDict.get('order'), differenceInPieces, orderDict.get('orderSheet'), prodProf, prodAv)

                #If the added number of skids leaves 0 skids remaining, then remove thhat object from inventory
                if prod.numberofskids + numSkids == 0:
                    prod.delete()
                else:
                    prod.numberofskids += Decimal(numSkids)
                    prod.save()
                return render(request, 'CWBDataApp/ProductInventoryUpdate.html', {'allProduct':allProduct, 'dataAcceptedMessage':"Successfully Added "+str(numSkids)+" Skids to "+ form['newcolour'] + " " + form['prodName']})
            else:
                return render(request, 'CWBDataApp/ProductInventoryUpdate.html', {'allProduct':allProduct, 'allColours':allColours, 'error_message':""+ form['newcolour'] + " " + form['prodName'] +" does not exist in inventory. If you wish to enter its data, press the Add Product tab"})
    return render(request, 'CWBDataApp/ProductInventoryUpdate.html', {'allProduct':allProduct, 'allColours':allColours})

###########################################################PRODUCT INVENTORY QUERY
def ProductInventoryQuery(request):

    allProduct = Productinventory.objects.all()

    if request.method == 'POST':
        form = request.POST

        if Productinventory.objects.filter(pk=form['prodId']).exists():
            prod = Productinventory.objects.get(pk=form['prodId'])
            return render(request, 'CWBDataApp/ProductInventoryQuery.html', {'allProduct':allProduct, 'product':prod})
        else:
            return render(request, 'CWBDataApp/ProductInventoryQuery.html', {'allProduct':allProduct, 'error_message':""+ form['colour'] + " " + form['prodName'] +" does not exist in inventory. If you wish to enter its data, press the link above"})
    return render(request, 'CWBDataApp/ProductInventoryQuery.html', {'allProduct':allProduct})

###########################################################Product Inventory Excel File Download
def ProductInventoryExcel(request):
    allProduct = Productinventory.objects.all()
    query = str(Productinventory.objects.all().query)
    df = pd.read_sql_query(query, connection)
    today = str(date.today())

    df.to_excel(r'./CWBDataApp/ProductInventory.xlsx', index=False)

    with open(r'./CWBDataApp/ProductInventory.xlsx', 'rb') as fh:
        response = HttpResponse(fh.read(), content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        response['Content-Disposition'] = 'attachment; filename=ProductInventory-'+today+'.xlsx'
    os.remove(r'./CWBDataApp/ProductInventory.xlsx')
    return response

    return render(request, 'CWBDataApp/ProductInventoryQuery.html', {'allProduct':allProduct})

###########################################################Product Inventory Excel File Download
def ProductInventoryPdf(request):
    allProduct = Productinventory.objects.all()
    query = str(Productinventory.objects.all().query)
    df = pd.read_sql_query(query, connection)
    fig, ax = plt.subplots(figsize=(12,4))
    today = date.today()

    ax.axis('tight')
    ax.axis('off')
    ax.set_title("Product Inventory " + str(today))
    the_table = ax.table(cellText=df.values, colLabels=df.columns, loc='center')

    pp = PdfPages(r'./CWBDataApp/ProductInventoryPdf.pdf')
    pp.savefig(fig, bbox_inches='tight')
    pp.close()

    return FileResponse(open(r'./CWBDataApp/ProductInventoryPdf.pdf', 'rb'), content_type='application/pdf')

    return render(request, 'CWBDataApp/ProductInventoryQuery.html', {'allProduct':allProduct})

###########################################################PRODUCT INVENTORY SHIPPED
def ProductInventoryShipped(request):

    allProduct = Productinventory.objects.all()

    if request.method == 'POST':
        form = request.POST

        if Productinventory.objects.filter(pk=form['prodName']).exists():
            prod = Productinventory.objects.get(pk=form['prodName'])

            PiecesRemaining = prod.numberofskids - Decimal(form['numSkids'])

            #Check if there's an order with this product
            Order = OrderFunc()
            orderDict = Order.FindOrder(prod.productname, prod.colour)
            #If an order exists with this product then update its inventorized pieces
            if orderDict:
                returnMessage = Order.updateOrderSent(orderDict.get('order'), Decimal(form['numSkids']), orderDict.get('orderSheet'))

                #If an error occured during the update of the order, then exit and send error message
                if returnMessage != '':
                    return render(request, 'CWBDataApp/ProductInventoryShipped.html', {'allProduct':allProduct, 'error_message':returnMessage})

            if PiecesRemaining <= 0:
                prod.delete()
                PiecesRemaining = 0
            else:
                prod.numberofskids = PiecesRemaining
                prod.save()
            return render(request, 'CWBDataApp/ProductInventoryShipped.html', {'allProduct':allProduct, 'dataAcceptedMessage':"Number of Skids Successfully Updated, "+str(PiecesRemaining)+" are remaining of "+str(prod.colour)+" "+str(prod.productname)+".",})
        else:
            return render(request, 'CWBDataApp/ProductInventoryShipped.html', {'allProduct':allProduct, 'error_message':"This product currently does not exist in inventory. If you wish to enter its data, press the link above"})
    return render(request, 'CWBDataApp/ProductInventoryShipped.html', {'allProduct':allProduct})

###########################################################MATERIAL INVENTORY
def MaterialInventory(request):

    allSuppliers= Materialcost.objects.all()

    if request.method == 'POST':
        form = request.POST

        try:
            Materialinventory.objects.get(pk=form['matName'])
            return render(request, 'CWBDataApp/MaterialInventory.html',{'allSuppliers':allSuppliers, 'error_message':"Material already exists in inventory, if you wish to update its data click the link above"})
        except:
            try:
                if form['numBoxes'] == '' or float(form['numBoxes']) <= 0 or float(form['price']) < 0:
                    return render(request, 'CWBDataApp/MaterialInventory.html',{'allSuppliers':allSuppliers, 'error_message':"Cannot enter zero for number of boxes or price"})

                sup_object = Materialcost.objects.get(pk=form['supplier'])
                new_material = Materialinventory(materialname=form['matName'].upper(),
                                                 supplier=sup_object,
                                                 numberofboxes=form['numBoxes'],
                                                 locations=form['location'],
                                                 premixed=form['premixed'],
                                                 priceperpound=(form['price']))
                new_material.save()
                return render(request, 'CWBDataApp/MaterialInventory.html', {'allSuppliers':allSuppliers, 'dataAcceptedMessage':"Material Successfully Added",})
            except:
                return render(request, 'CWBDataApp/MaterialInventory.html',{'allSuppliers':allSuppliers, 'error_message':"Supplier does not exist"})


    return render(request, 'CWBDataApp/MaterialInventory.html', {'allSuppliers':allSuppliers})

###########################################################MATERIAL INVENTORY QUERY
def MaterialInventoryQuery(request):

    allMaterial = Materialinventory.objects.all()

    if request.method == 'POST':
        form = request.POST
        try:
            if form['matName'].upper() == 'ALL':
                return render(request, 'CWBDataApp/MaterialInventoryQuery.html', {'allMaterial':allMaterial, 'all_materials':all_materials})


            mat_object =  Materialinventory.objects.filter(pk=form['matName'].upper())
            return render(request, 'CWBDataApp/MaterialInventoryQuery.html', {'allMaterial':allMaterial, 'all_materials':mat_object})
        except:
            return render(request, 'CWBDataApp/MaterialInventoryQuery.html', {'allMaterial':allMaterial, 'error_message':"This material currently does not exist in inventory. If you wish to enter its data, press the link above"})

    return render(request, 'CWBDataApp/MaterialInventoryQuery.html', {'allMaterial':allMaterial})

###########################################################Material Inventory Excel File Download
def MaterialInventoryExcel(request):
    allMaterial = Materialinventory.objects.all()
    query = str(Materialinventory.objects.all().query)
    df = pd.read_sql_query(query, connection)
    today = str(date.today())

    df.to_excel(r'./CWBDataApp/MaterialInventory.xlsx', index=False)

    with open(r'./CWBDataApp/MaterialInventory.xlsx', 'rb') as fh:
        response = HttpResponse(fh.read(), content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        response['Content-Disposition'] = 'attachment; filename=MaterialInventory-'+today+'.xlsx'
    os.remove(r'./CWBDataApp/MaterialInventory.xlsx')
    return response

    return render(request, 'CWBDataApp/MaterialInventoryQuery.html', {'allMaterial':allMaterial})

###########################################################Material Inventory PDF Print Download
def MaterialInventoryPdf(request):
    allMaterial = Materialinventory.objects.all()
    query = str(Materialinventory.objects.all().query)
    df = pd.read_sql_query(query, connection)
    fig, ax = plt.subplots()
    today = date.today()


    ax.axis('off')
    ax.axis('tight')
    ax.set_title("Material Inventory " + str(today) + "\n\n\n", fontsize="20")
    the_table = ax.table(cellText=df.values, colLabels=df.columns, loc='center')
    fig.tight_layout()

    pp = PdfPages(r'./CWBDataApp/MaterialInventoryPdf.pdf')
    pp.savefig(fig, bbox_inches='tight')
    pp.close()


    return FileResponse(open(r'./CWBDataApp/MaterialInventoryPdf.pdf', 'rb'), content_type='application/pdf')

    return render(request, 'CWBDataApp/MaterialInventoryQuery.html', {'allMaterial':allMaterial})


###########################################################MATERIAL INVENTORY UPDATE
def MaterialInventoryUpdate(request):

    allMaterials = Materialinventory.objects.all()

    if request.method == 'POST':
        form = request.POST
        try:
            if form['matName2'] != 'None':
                mat_object = Materialinventory.objects.get(pk=form['matName2'].upper())
                return render(request, 'CWBDataApp/MaterialInventoryUpdate.html', {'allMaterials':allMaterials, 'Forms':True, 'Material':mat_object})
            elif Materialinventory.objects.get(pk=form['matName'].upper()):
                mat_object = Materialinventory.objects.get(pk=form['matName'].upper())
                return render(request, 'CWBDataApp/MaterialInventoryUpdate.html', {'allMaterials':allMaterials, 'Forms':True, 'Material':mat_object})
            else:
                return render(request, 'CWBDataApp/MaterialInventoryUpdate.html', {'allMaterials':allMaterials, 'error_message':"Material does not exist in inventory, check your spelling"})

        except:
            return render(request, 'CWBDataApp/MaterialInventoryUpdate.html', {'allMaterials':allMaterials, 'error_message':"This material currently does not exist in inventory. If you wish to enter its data, press the link above"})


    return render(request, 'CWBDataApp/MaterialInventoryUpdate.html', {'allMaterials':allMaterials})

###########################################################MATERIAL INVENTORY UPDATE NUMBER of SKIDS
def MaterialInventoryUpdateNumSkids(request):

    allMaterials = Materialinventory.objects.all()

    if request.method == 'POST':
        form = request.POST
        try:
            if form['numBoxes'] == '' or float(form['numBoxes']) < 0 or float(form['price']) < 0:
                return render(request, 'CWBDataApp/MaterialInventoryUpdate.html',{'allMaterials':allMaterials, 'error_message':"Cannot enter less than zero for number of boxes"})
            elif float(form['numBoxes']) == 0:
                mat_object = Materialinventory.objects.get(pk=form['matName'].upper())
                mat_object.delete()
                return render(request, 'CWBDataApp/MaterialInventoryUpdate.html', {'allMaterials':allMaterials, 'dataAcceptedMessage':"Material Was Successfully Deleted Since Zero Boxes Remained"})

            mat_object = Materialinventory.objects.get(pk=form['matName'].upper())
            mat_object.numberofboxes = form['numBoxes']
            mat_object.locations = form['location']
            mat_object.priceperpound = form['price']
            mat_object.save()
            return render(request, 'CWBDataApp/MaterialInventoryUpdate.html', {'allMaterials':allMaterials, 'dataAcceptedMessage':"Material Updated"})
        except:
            return render(request, 'CWBDataApp/MaterialInventoryUpdate.html', {'allMaterials':allMaterials, 'error_message':"Could not update material, try again"})


    return render(request, 'CWBDataApp/MaterialInventoryUpdate.html', {'allMaterials':allMaterials})

###########################################################ORDER SHEET MACHINE 1
def OrderSheetsMachine1(request):
    allColours = Colour.objects.all()
    allProfiles= Productprofiles.objects.all()

    if request.method == 'POST':
        form = request.POST
        orderModel = globals()["Ordersheetmachine1"]
        message = AddOrder(form, orderModel)

        if message['Error']:
            return render(request, 'CWBDataApp/OrderSheetsMachine1.html', {'allColours':allColours, 'allProfiles':allProfiles, 'error_message':message['error_message']})

        return render(request, 'CWBDataApp/OrderSheetsMachine1.html', {'allColours':allColours, 'allProfiles':allProfiles, 'dataAcceptedMessage':"Order Successfully Added"})
    return render(request, 'CWBDataApp/OrderSheetsMachine1.html', {'allColours':allColours, 'allProfiles':allProfiles})


###########################################################ORDER SHEET MACHINE 2
def OrderSheetsMachine2(request):
    allColours = Colour.objects.all()
    allProfiles= Productprofiles.objects.all()

    if request.method == 'POST':
        form = request.POST
        orderModel = globals()["Ordersheetmachine2"]
        message = AddOrder(form, orderModel)

        if message['Error']:
            return render(request, 'CWBDataApp/OrderSheetsMachine2.html', {'allColours':allColours, 'allProfiles':allProfiles, 'error_message':message['error_message']})

        return render(request, 'CWBDataApp/OrderSheetsMachine2.html', {'allColours':allColours, 'allProfiles':allProfiles, 'dataAcceptedMessage':"Order Successfully Added"})
    return render(request, 'CWBDataApp/OrderSheetsMachine2.html', {'allColours':allColours, 'allProfiles':allProfiles})

###########################################################ORDER SHEET MACHINE 3
def OrderSheetsMachine3(request):
    allColours = Colour.objects.all()
    allProfiles= Productprofiles.objects.all()

    if request.method == 'POST':
        form = request.POST
        orderModel = globals()["Ordersheetmachine3"]
        message = AddOrder(form, orderModel)

        if message['Error']:
            return render(request, 'CWBDataApp/OrderSheetsMachine3.html', {'allColours':allColours, 'allProfiles':allProfiles, 'error_message':message['error_message']})

        return render(request, 'CWBDataApp/OrderSheetsMachine3.html', {'allColours':allColours, 'allProfiles':allProfiles, 'dataAcceptedMessage':"Order Successfully Added"})
    return render(request, 'CWBDataApp/OrderSheetsMachine3.html', {'allColours':allColours, 'allProfiles':allProfiles})

###########################################################Helper Function Used to Add Orders
def AddOrder(form, Ordersheet):
    pcsinventorized = 0
    productProfile = Productprofiles.objects.get(pk=form['prodName'])
    #If the product exists in inventory and it isnt already in the ordersheets, then add the inventory to this order
    if Productinventory.objects.filter(productname=form['prodName'], colour=form['colour']).exists() and Ordersheet.objects.filter(boardprofile=form['prodName'], colour=form['colour']).exists() == False:
        product = Productinventory.objects.filter(productname=form['prodName'], colour=form['colour']).first()
        pcsinventorized = int(product.numberofskids) * productProfile.pcsperskid


    avgSkidsDay = 1
    #If no average skid number exist return error
    if Profileaverages.objects.filter(pk=form['prodName']).exists():
        productAverage = Profileaverages.objects.get(pk=form['prodName'])
        avgSkidsDay = productAverage.averageskidsperday
    else:
        return {'Error' : True, 'error_message':'No average Skids per day exists for this profile. Please enter it through general tab'}

    #Initialize the orders variables
    pcsremaining = int(form['pcs']) - int(form['pcsSent']) - pcsinventorized
    skidsremaining = pcsremaining / int(productProfile.pcsperskid)
    lengthofrunindays = Decimal(skidsremaining) / Decimal(productAverage.averageskidsperday)
    priority = 0
    startdate = date.today()
    enddate = addDate(startdate, math.ceil(lengthofrunindays))
    duedate = enddate
    ponumber = form['ponumber']

    #If there are no orders then priority should be 1
    if Ordersheet.objects.all() == None:
        priority = 1
    else:
        if form['priority'] == '' :
            #If no priority is given, grab the lowest priority entry and give this order a lower priority
            lastOrder = Ordersheet.objects.all().order_by('-priorityrank').first()
            priority = lastOrder.priorityrank + 1
            startdate = lastOrder.projectedenddate
            enddate = addDate(startdate, math.ceil(lengthofrunindays))


        else:
            #If a priority is given, then find where it belongs and shuffle lower orders around
            lessPriority = Ordersheet.objects.filter(priorityrank__gt = (int(form['priority'])-1)).order_by('priorityrank')
            if int(form['priority']) == 1:
                startdate = date.today()
            else:
                higherpriority = Ordersheet.objects.filter(priorityrank=(int(form['priority'])-1)).first()
                startdate = higherpriority.projectedenddate
            enddate = addDate(startdate, math.ceil(lengthofrunindays))
            Order = OrderFunc()
            #Call Function to update the priority of lower orders and their respective dates
            Order.deprioritize(form['priority'], enddate, lessPriority)
            priority = form['priority']

    #If no due date is given use the orders end date
    if form['dueDate'] == '':
        duedate = enddate
    else:
        duedate = form['dueDate']

    #Create and save the new order object
    newOrder=Ordersheet( customerponumber=ponumber,
                         projectedstartdate=startdate,
                         projectedenddate=enddate,
                         duedate=duedate,
                         lengthofrunindays=lengthofrunindays,
                         priorityrank=priority,
                         boardprofile=form['prodName'],
                         colour=form['colour'],
                         skidsremaining=skidsremaining,
                         pcs=form['pcs'],
                         pcssent=form['pcsSent'],
                         pcsremaining=pcsremaining,
                         customer=form['customer'],
                         qualitynotes=form['qualitynotes'],
                         pcsinventorized=pcsinventorized)
    newOrder.save()
    return {'Error': False}

###########################################################Update DD Orders
def UpdateDDOrders(request):

    if request.method == 'POST':
        form = request.POST
        allProfiles= Productprofiles.objects.all()
        allColours = Colour.objects.all()
        orderModel = globals()["Ordersheetmachine" + form['machine']]
        #Call DDOrder Function to grab DD Orders from Google sheets and clean data
        order = OrdersDD()
        orders = order.getDDOrders(allProfiles, allColours, form['machine'])



        for key, value in orders.items():

            #Split Key into product and colour
            prodString = key.split('_')
            product = prodString[0]
            colour = prodString[1]

            pcssent = value[1]
            pcs = value[0]

            productProfile = Productprofiles.objects.get(pk=product)
            productAverage = Profileaverages.objects.get(pk=product)

            #if an order already exists for this profile then update its values if needed
            if orderModel.objects.filter(boardprofile=product, colour=colour).exists():
                currentOrder = orderModel.objects.get(boardprofile=product, colour=colour)

                #If the values of total pieces or total sent have changed then update the order
                if currentOrder.pcssent != pcssent or currentOrder.pcs != pcs:

                    currentOrder.pcssent = pcssent
                    currentOrder.pcs = pcs
                    currentOrder.pcsremaining = pcs - pcssent - currentOrder.pcsinventorized
                    currentOrder.skidsremaining = currentOrder.pcsremaining / int(productProfile.pcsperskid)
                    currentOrder.lengthofrunindays = Decimal(currentOrder.skidsremaining) / Decimal(productAverage.averageskidsperday)
                    currentOrder.projectedenddate = addDate(currentOrder.projectedstartdate, math.ceil(currentOrder.lengthofrunindays))
                    currentOrder.save()
                    #Grab all orders with lower priority
                    lowerOrders = orderModel.objects.filter(priorityrank__gt=form['priority']).order_by('priorityrank')
                    endDate = order.projectedenddate
                    for item in lowerOrders:
                        item.projectedstartdate = endDate
                        item.projectedenddate = addDate(endDate, math.ceil(item.lengthofrunindays))
                        item.save()
                        endDate = item.projectedenddate
            else:

                pcsinventorized = 0

                #Check if this product has any inventory
                if Productinventory.objects.filter(productname=product, colour=colour).exists() and orderModel.objects.filter(boardprofile=product, colour=colour).exists() == False:
                    product = Productinventory.objects.filter(productname=product, colour=colour).first()
                    pcsinventorized = int(product.numberofskids)

                pcsremaining = pcs - pcssent - pcsinventorized
                skidsremaining = pcsremaining / int(productProfile.pcsperskid)
                lengthofrunindays = Decimal(skidsremaining) / Decimal(productAverage.averageskidsperday)
                priority = 1
                startdate = date.today()
                enddate = addDate(startdate, math.ceil(lengthofrunindays))

                #If Orders exist, then give this order the lowest priority
                if orderModel.objects.all().order_by('-priorityrank').exists():
                    lastOrder = orderModel.objects.all().order_by('-priorityrank').first()
                    priority = lastOrder.priorityrank + 1
                    startdate = lastOrder.projectedenddate
                    enddate = addDate(startdate, math.ceil(lengthofrunindays))
                duedate = enddate

                newOrder= orderModel( customerponumber='',
                                      projectedstartdate=startdate,
                                      projectedenddate=enddate,
                                      duedate=duedate,
                                      lengthofrunindays=lengthofrunindays,
                                      priorityrank=priority,
                                      boardprofile=product,
                                      colour=colour,
                                      skidsremaining=skidsremaining,
                                      pcs=pcs,
                                      pcssent=pcssent,
                                      pcsremaining=pcsremaining,
                                      customer='DD',
                                      qualitynotes='',
                                      pcsinventorized=pcsinventorized)
                newOrder.save()
        return render(request, 'CWBDataApp/DDOrders.html', {'dataAcceptedMessage': 'Successfully Updated DDOrders', 'numberOfMachines':range(1, numberOfMachines+1)})
    return render(request, 'CWBDataApp/DDOrders.html', {'numberOfMachines':range(1, numberOfMachines+1)})

###########################################################ORDER SHEET Helpers
def ViewOrders(request):

    if request.method == 'POST':
        form = request.POST
        dict = {}

        if form['machine'] == 'all':
            dict['machine1'] = Ordersheetmachine1.objects.all().order_by('priorityrank')
            dict['machine2'] = Ordersheetmachine2.objects.all().order_by('priorityrank')
            dict['machine3'] = Ordersheetmachine3.objects.all().order_by('priorityrank')
        else:
            orderModel = globals()["Ordersheetmachine" + form['machine']]
            if orderModel.objects.all():
                dict['machine'+ form['machine']] = orderModel.objects.all().order_by('priorityrank')

        if not dict:
            return render(request, 'CWBDataApp/ViewOrders.html', {'numberOfMachines':range(1, numberOfMachines+1), 'error_message': "There are Currently No Orders For This Machine"})
        else:
            return render(request, 'CWBDataApp/ViewOrders.html', {'numberOfMachines':range(1, numberOfMachines+1), 'allOrders':dict.values()})

    return render(request, 'CWBDataApp/ViewOrders.html', {'numberOfMachines':range(1, numberOfMachines+1)})


###########################################################Order Excel File Download
def OrdersExcel(request):

    writer = pd.ExcelWriter(r'./CWBDataApp/OrderSheetMachine.xlsx')
    today = str(date.today())
    for i in range(1, numberOfMachines+1):
        orderModel = globals()["Ordersheetmachine" + str(i)]
        query = str(orderModel.objects.all().query)
        df = pd.read_sql_query(query, connection)
        df.to_excel(writer, sheet_name= 'Machine'+str(i) ,index=False)
    writer.close()

    with open(r'./CWBDataApp/OrderSheetMachine.xlsx', 'rb') as fh:
        response = HttpResponse(fh.read(), content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        response['Content-Disposition'] = 'attachment; filename=OrderSheetMachine-'+today+'.xlsx'
    os.remove(r'./CWBDataApp/OrderSheetMachine.xlsx')
    return response

    return render(request, 'CWBDataApp/ViewOrders.html', {'numberOfMachines':range(1, numberOfMachines+1)})


###########################################################UPDATE ORDERS
def UpdateOrders(request):

    if request.method == 'POST':
        form = request.POST
        orders = None

        orderModel = globals()["Ordersheetmachine" + form['machine']]
        if orderModel.objects.all():
            orders = orderModel.objects.all().order_by('priorityrank')

        if orders == None:
            return render(request, 'CWBDataApp/UpdateOrders.html', {'numberOfMachines':range(1, numberOfMachines+1), 'error_message':"No Orders Exist For This Machine"})
        else:
            return render(request, 'CWBDataApp/UpdateOrders.html', {'numberOfMachines':range(1, numberOfMachines+1), 'orders':orders, 'machine':form['machine']})

    return render(request, 'CWBDataApp/UpdateOrders.html', {'numberOfMachines':range(1, numberOfMachines+1)})

###########################################################GET SPECIFIC ORDER TO UPDATE
def GetOrder(request):

    if request.method == 'POST':
        form = request.POST

        try:
            orderModel = globals()["Ordersheetmachine" + form['machine']]
            order = orderModel.objects.get(priorityrank=form['orderrank'])
            return render(request, 'CWBDataApp/UpdateOrders.html', {'numberOfMachines':range(1, numberOfMachines+1), 'orderChange':order, 'machine':form['machine']})
        except:
            return render(request, 'CWBDataApp/UpdateOrders.html', {'numberOfMachines':range(1, numberOfMachines+1), 'error_message':"Could not grab orders"})
    return render(request, 'CWBDataApp/UpdateOrders.html', {'numberOfMachines':range(1, numberOfMachines+1)})


###########################################################REMOVE SPECIFIC ORDER
def RemoveOrder(request):

    if request.method == 'POST':
        form = request.POST

        try:
            orderModel = globals()["Ordersheetmachine" + form['machine']]
            order = orderModel.objects.get(priorityrank=form['orderrank'])
            orderFunc = OrderFunc()
            #Call Function to update the priority of lower orders and their respective dates
            orderFunc.deletedOrder(order, orderModel.objects.all())
            order.delete()
            return render(request, 'CWBDataApp/UpdateOrders.html', {'dataAcceptedMessage':'Order Successfully Removed', 'numberOfMachines':range(1, numberOfMachines+1)})
        except:
            return render(request, 'CWBDataApp/UpdateOrders.html', {'numberOfMachines':range(1, numberOfMachines+1), 'error_message':"Could not get Order"})
    return render(request, 'CWBDataApp/UpdateOrders.html', {'numberOfMachines':range(1, numberOfMachines+1)})

###########################################################Change ORDER
def ChangeOrder(request):

    if request.method == 'POST':
        form = request.POST


        orderModel = globals()["Ordersheetmachine" + form['machine']]
        order = orderModel.objects.get(pk=form['orderPK'])
        order.customerponumber = form['ponumber']
        order.customer = form['customer']
        order.qualitynotes = form['qualitynotes']

        #If the total number of pieces or the number of pieces sent changes, then update all values and shuffle dates around
        if order.pcs != int(form['pcs']) or order.pcssent != int(form['pcsSent']):

            if order.pcs != int(form['pcs']):
                order.pcs = int(form['pcs'])

            if order.pcssent != int(form['pcsSent']):
                order.pcssent = int(form['pcsSent'])

            order.save()
            order.pcsremaining = order.pcs - order.pcssent - order.pcsinventorized
            product = Productprofiles.objects.get(pk=order.boardprofile)
            productAverage = Profileaverages.objects.get(pk=order.boardprofile)
            order.skidsremaining = order.pcsremaining / int(product.pcsperskid)
            order.lengthofrunindays = Decimal(order.skidsremaining) / Decimal(productAverage.averageskidsperday)
            order.projectedenddate = addDate(order.projectedstartdate, math.ceil(order.lengthofrunindays))
            order.save()
            #Grab all orders with lower priority
            lowerOrders = orderModel.objects.filter(priorityrank__gt=form['priority']).order_by('priorityrank')
            endDate = order.projectedenddate
            for item in lowerOrders:
                item.projectedstartdate = endDate
                item.projectedenddate = addDate(endDate, math.ceil(item.lengthofrunindays))
                item.save()
                endDate = item.projectedenddate


        #If priority was changed then reorder all the order where necessary
        #If new priority is higher than the current, shuffle all orders down
        if int(form['priority']) < order.priorityrank:
            startdate = None
            lessPriority = orderModel.objects.filter(priorityrank__gte = (int(form['priority']))).order_by('priorityrank')
            if int(form['priority']) == 1:
                startdate = date.today()
            else:
                higherpriority = orderModel.objects.filter(priorityrank=(int(form['priority'])-1)).first()
                startdate = higherpriority.projectedenddate
            enddate = addDate(startdate, math.ceil(order.lengthofrunindays))

            #Loop to change priority and dates of order the are under the new priority
            for entry in lessPriority:
                if entry.priorityrank == form['priority']:
                    entry.projectedstartdate = endDate
                    entry.projectedenddate = endDate + timedelta(days=math.ceil(entry.lengthofrunindays))
                else:
                    higherpriority = orderModel.objects.filter(priorityrank=entry.priorityrank).first()
                    entry.projectedstartdate = higherpriority.projectedenddate
                    entry.projectedenddate =  higherpriority.projectedenddate + timedelta(days=math.ceil(entry.lengthofrunindays))
                #only change priority of orders in which the priority was shifted
                if entry.priorityrank < order.priorityrank:
                    entry.priorityrank += 1
                entry.save()

            order.priorityrank = form['priority']
            order.projectedstartdate = startdate
            order.projectedenddate = enddate
            order.save()
        elif int(form['priority']) > order.priorityrank:
            startdate = None
            lessPriority = orderModel.objects.filter(priorityrank__gt = order.priorityrank ).order_by('priorityrank')

            if order.priorityrank == 1:
                startdate = date.today()
            else:
                higherpriority = orderModel.objects.filter(priorityrank=order.priorityrank).first()
                startdate = higherpriority.projectedenddate

            for entry in lessPriority:
                if entry.priorityrank == int(order.priorityrank) + 1:
                    entry.projectedstartdate = startdate
                    entry.projectedenddate = addDate(startdate, math.ceil(entry.lengthofrunindays))
                elif entry.priorityrank > order.priorityrank+1 and entry.priorityrank <= int(form['priority']):
                    higherpriority = orderModel.objects.filter(priorityrank=entry.priorityrank-2).first()
                    entry.projectedstartdate = higherpriority.projectedenddate
                    entry.projectedenddate =  addDate(higherpriority.projectedenddate, math.ceil(entry.lengthofrunindays))
                elif entry.priorityrank > form['priority']:
                    higherpriority = orderModel.objects.filter(priorityrank=entry.priorityrank-1).first()
                    entry.projectedstartdate = higherpriority.projectedenddate
                    entry.projectedenddate =  addDate(higherpriority.projectedenddate, math.ceil(entry.lengthofrunindays))

                if entry.priorityrank == int(form['priority']):
                    order.priorityrank = form['priority']
                    order.projectedstartdate = entry.projectedenddate
                    order.projectedenddate = addDate(entry.projectedenddate, math.ceil(order.lengthofrunindays))
                    order.save()
                #only change priority of orders in which the priority was shifted
                if entry.priorityrank <= int(form['priority']):
                    entry.priorityrank -= 1
                entry.save()

        order.save()

        return render(request, 'CWBDataApp/UpdateOrders.html', {'numberOfMachines':range(1, numberOfMachines+1), 'dataAcceptedMessage':"Order updated"})
    return render(request, 'CWBDataApp/UpdateOrders.html', {'numberOfMachines':range(1, numberOfMachines+1)})


###########################################################PIC & SUM
def PicSum(request):
    allEmployees = Employees.objects.all()

    if request.method == 'POST':
        form = picsumForm(request.POST, request.FILES)

        if form.is_valid():
            form.save()
            return render(request, 'CWBDataApp/PicSum.html', {'allEmployees':allEmployees, 'dateToday':date.today().isoformat(), 'dataAcceptedMessage':"Pic&Sum Successfully Uploaded"})
        else:
            print(form.errors)
            return render(request, 'CWBDataApp/PicSum.html', {'allEmployees':allEmployees, 'dateToday':date.today().isoformat(), 'error_message':"Something went wrong could not upload. Check to make sure everything is entered correctly"})


    return render(request, 'CWBDataApp/PicSum.html', {'allEmployees':allEmployees, 'dateToday':date.today().isoformat()})

###########################################################PIC & SUM
def PicSumForms(request):

    if request.method == 'POST':
        form = request.POST

        query = str(picsum.objects.all().query)
        df = pd.read_sql_query(query, connection)
        today = str(date.today())
        df.drop(columns=['image', 'description'], inplace=True)
        df.rename(columns={'title': 'Test Name', 'testdate':'Test Date', 'supervisor':'Supervisor (Not a temp)', 'machineoperator':'Machine Operator', 'temp1':'Is the Machine Operator a Temp', 'mixer':'Mixer', 'temp2':'Is the Mixer a temp'}, inplace=True)

        df['Test Date'] = pd.to_datetime(df['Test Date'], format='%Y-%m-%d')

        if form['startdate'] != '':
            startdate = datetime.strptime(form['startdate'], '%Y-%m-%d')
            df = df.loc[(df['Test Date'] >= startdate)]

        if form['enddate'] != '':
            enddate = datetime.strptime(form['enddate'], '%Y-%m-%d')
            df = df.loc[(df['Test Date'] <= enddate)]

        df.to_excel(r'./CWBDataApp/picsum.xlsx', index=False)

        with open(r'./CWBDataApp/picsum.xlsx', 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            response['Content-Disposition'] = 'attachment; filename=PicSum-'+today+'.xlsx'
        os.remove(r'./CWBDataApp/picsum.xlsx')
        return response


    return render(request, 'CWBDataApp/PicSumForms.html')

###########################################################PIC & SUM Doc generator
def PicSumDoc(request):

    if request.method == 'POST':
        form = request.POST

        #Check if end and start dates are given
        if form['startdate'] != '':
            startdate = datetime.strptime(form['startdate'], '%Y-%m-%d').date()

            if form['enddate'] != '':
                enddate = datetime.strptime(form['enddate'], '%Y-%m-%d').date()
                allPics = picsum.objects.filter(testdate__gte=startdate, testdate__lte=enddate)
            else:
                allPics = picsum.objects.filter(testdate__gte=startdate)
        elif form['enddate'] != '':
            enddate = datetime.strptime(form['enddate'], '%Y-%m-%d').date()
            allPics = picsum.objects.filter(testdate__lte=enddate)
        else:
            allPics = picsum.objects.all()


        myDoc = docx.Document()

        #Create doc with all pic&PicSum
        #Include title, testdate, decription and image
        for pic in allPics:
            header = myDoc.add_paragraph().add_run(pic.title)
            header.font.underline = True
            header.font.size = docx.shared.Pt(18)
            header.font.bold = True
            myDoc.add_heading(str(pic.testdate), 2)

            myDoc.add_paragraph(pic.description)
            myDoc.add_picture(pic.image, width=docx.shared.Inches(4), height=docx.shared.Inches(3))
            myDoc.save(r'./CWBDataApp/Pic&SumDoc.docx')

        with open(r'./CWBDataApp/Pic&SumDoc.docx', 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            response['Content-Disposition'] = 'attachment; filename=PicSumDoc.docx'
        os.remove(r'./CWBDataApp/Pic&SumDoc.docx')
        return response


    return render(request, 'CWBDataApp/PicSumForms.html')


###########################################################HELP
def help(request):
    return render(request, 'CWBDataApp/help.html')

###########################################################Add Employee
def AddEmployee(request):
    allEmployees = Employees.objects.all()

    if request.method == 'POST':
        form = request.POST

        try:
            employee = Employees.objects.get(pk=form['employeeName'])
            return render(request, 'CWBDataApp/AddEmployee.html', {'allEmployees':allEmployees, 'error_message':"Employee already exists"})
        except:
            new_employee = Employees(employeename=form['employeeName'])
            new_employee.save()
            return render(request, 'CWBDataApp/AddEmployee.html', {'allEmployees':allEmployees,'dataAcceptedMessage':"Employee Successfully Added"})
    return render(request, 'CWBDataApp/AddEmployee.html', {'allEmployees':allEmployees})

###########################################################Remove Employee
def RemoveEmployee(request):
    allEmployees = Employees.objects.all()

    if request.method == 'POST':
        form = request.POST

        try:
            employee = Employees.objects.get(pk=form['employeeName'])
            employee.delete()
            return render(request, 'CWBDataApp/AddEmployee.html', {'allEmployees':allEmployees, 'dataAcceptedMessage':"Employee Successfully Removed"})
        except:
            return render(request, 'CWBDataApp/AddEmployee.html', {'allEmployees':allEmployees, 'error_message':"Employee Cannot Be Removed Because They Do Not Exist"})
    return render(request, 'CWBDataApp/AddEmployee.html', {'allEmployees':allEmployees})

###########################################################Add Board Profile
def AddBoardProfile(request):
    allProfiles = Productprofiles.objects.all()

    if request.method == 'POST':
        form = request.POST

        try:
            profile = Productprofiles.objects.get(pk=form['profile'])
            return render(request, 'CWBDataApp/AddBoardProfile.html', {'allProfiles':allProfiles, 'error_message':"Profile already exists"})
        except:
            pcs=0
            if form['pcsperskid'] != '':
                pcs = int(form['pcsperskid'])
            new_profile = Productprofiles(productname=form['profile'], pcsperskid=pcs)
            new_profile.save()
            return render(request, 'CWBDataApp/AddBoardProfile.html', {'allProfiles':allProfiles, 'dataAcceptedMessage':"Profile Successfully Added"})
    return render(request, 'CWBDataApp/AddBoardProfile.html', {'allProfiles':allProfiles})

###########################################################Add Board Profile
def RemoveBoardProfile(request):
    allProfiles = Productprofiles.objects.all()

    if request.method == 'POST':
        form = request.POST

        try:
            if form['profile'] == '':
                return render(request, 'CWBDataApp/AddBoardProfile.html', {'allProfiles':allProfiles, 'error_message':"Please enter a profile"})
            profile = Productprofiles.objects.get(pk=form['profile'])
            profile.delete()
            return render(request, 'CWBDataApp/AddBoardProfile.html', {'allProfiles':allProfiles, 'dataAcceptedMessage':"Profile Successfully Removed"} )
        except:
            return render(request, 'CWBDataApp/AddBoardProfile.html', {'allProfiles':allProfiles, 'error_message':"Profile cannot be removed because it does not exist"})
    return render(request, 'CWBDataApp/AddBoardProfile.html', {'allProfiles':allProfiles})


###########################################################Add Colour
def AddColour(request):

    if request.method == 'POST':
        form = request.POST

        try:
            colour = Colour.objects.get(pk=form['colour'])
            return render(request, 'CWBDataApp/AddColour.html', {'error_message':"Colour already exists"})
        except:
            new_colour = Colour(colour=form['colour'])
            new_colour.save()
            return render(request, 'CWBDataApp/AddColour.html', {'dataAcceptedMessage':"Colour Successfully Added"})
    return render(request, 'CWBDataApp/AddColour.html')

###########################################################Add Colour
def RemoveColour(request):

    if request.method == 'POST':
        form = request.POST

        try:
            colour = Colour.objects.get(pk=form['colour'])
            colour.delete()
            return render(request, 'CWBDataApp/AddColour.html', {'dataAcceptedMessage':"Colour Successfully Deleted"})
        except:
            return render(request, 'CWBDataApp/AddColour.html', {'error_message':"Colour could not be removed because it does not exist"})
    return render(request, 'CWBDataApp/AddColour.html')

###########################################################Add Supplier
def AddSupplier(request):

    if request.method == 'POST':
        form = request.POST

        try:
            supplier = Materialcost.objects.get(pk=form['supplier'])
            return render(request, 'CWBDataApp/AddSupplier.html', {'error_message':"Supplier already exists"})
        except:
            new_supplier = Materialcost(supplier=form['supplier'],
                                        costperpound=form['price']
                                       )
            new_supplier.save()
            return render(request, 'CWBDataApp/AddSupplier.html', {'dataAcceptedMessage':"Supplier Successfully Added"})

    return render(request, 'CWBDataApp/AddSupplier.html')

###########################################################Remove Supplier
def RemoveSupplier(request):

    if request.method == 'POST':
        form = request.POST

        try:
            supplier = Materialcost.objects.get(pk=form['supplier'])
            supplier.delete()
            return render(request, 'CWBDataApp/AddSupplier.html', {'dataAcceptedMessage':"Supplier Successfully Removed"})
        except:
            return render(request, 'CWBDataApp/AddSupplier.html', {'error_message':"Supplier Cannot Be Removed Because It Does Not Exist"})
    return render(request, 'CWBDataApp/AddSupplier.html')

###########################################################Update Supplier
def UpdateSupplier(request):
    allSuppliers = Materialcost.objects.all()

    if request.method == 'POST':
        form = request.POST

        try:
            supplier = Materialcost.objects.get(pk=form['supplier'])
            supplier.costperpound = form['price']
            supplier.save()
            return render(request, 'CWBDataApp/UpdateSupplier.html', {'allSuppliers': allSuppliers, 'dataAcceptedMessage':"Supplier Successfully Updated"})
        except:
            return render(request, 'CWBDataApp/UpdateSupplier.html', {'allSuppliers':allSuppliers, 'error_message':'Supplier does not exist'})

    return render(request, 'CWBDataApp/UpdateSupplier.html', {'allSuppliers': allSuppliers})

###########################################################ADD PROFILE AVERAGE
def AddProfileAverage(request):

    if request.method == 'POST':
        form = request.POST

        try:
            profile = Profileaverages.objects.get(pk=form['profile'])
            return render(request, 'CWBDataApp/AddProfileAverage.html', {'error_message':"Profile already exists. If you wish to update this profile use, the UpdateProfileAverage tab under help"})
        except:
            new_average = Profileaverages(productname=form['profile'],
                                          averageskidsperday=form['average']
                                       )
            new_average.save()
            return render(request, 'CWBDataApp/AddProfileAverage.html', {'dataAcceptedMessage':"Profile Average Successfully Added"})


    return render(request, 'CWBDataApp/AddProfileAverage.html')

###########################################################UPDATE PROFILE AVERAGE
def UpdateProfileAverage(request):
    allAverages= Profileaverages.objects.all()

    if request.method == 'POST':
        form = request.POST

        try:
            profile = Profileaverages.objects.get(pk=form['profile'])
            profile.averageskidsperday = form['average']
            profile.save()
            return render(request, 'CWBDataApp/UpdateProfileAverage.html', {'allAverages':allAverages, 'dataAcceptedMessage':"Profile Average Successfully Updated"})
        except:
            return render(request, 'CWBDataApp/UpdateProfileAverage.html', {'allAverages':allAverages, 'error_message':"Profile could not be updated"})

    return render(request, 'CWBDataApp/UpdateProfileAverage.html', {'allAverages':allAverages})

###########################################################Add CIS Email subject
def AddSubject(request):
    allSubjects = cisEmailSubject.objects.all()

    if request.method == 'POST':
        form = request.POST

        try:
            subject = cisEmailSubject.objects.get(pk=form['subject'])
            return render(request, 'CWBDataApp/AddSubject.html', {'error_message':"Subject already exists", 'allSubjects':allSubjects})
        except:
            new_subject = cisEmailSubject(subject=form['subject'])
            new_subject.save()
            return render(request, 'CWBDataApp/AddSubject.html', {'dataAcceptedMessage':"Subject Successfully Added", 'allSubjects':allSubjects})
    return render(request, 'CWBDataApp/AddSubject.html', {'allSubjects':allSubjects})

###########################################################Remove Subject
def RemoveSubject(request):
    allSubjects = cisEmailSubject.objects.all()

    if request.method == 'POST':
        form = request.POST

        try:
            subject = cisEmailSubject.objects.get(pk=form['subject'])
            subject.delete()
            return render(request, 'CWBDataApp/AddSubject.html', {'dataAcceptedMessage':"Subject Successfully Deleted", 'allSubjects':allSubjects})
        except:
            return render(request, 'CWBDataApp/AddSubject.html', {'error_message':"Subjectt could not be removed because it does not exist", 'allSubjects':allSubjects})
    return render(request, 'CWBDataApp/AddSubject.html', {'allSubjects':allSubjects})

###########################################################REPORT ERRORS AND SUGGESTIONS
def Report(request):

    allSubjects = cisEmailSubject.objects.all()

    if request.method == 'POST':
        form = request.POST

        message = sendEmail(form['subject'], form['message'])

        if message['state'] == True:
            return render(request, 'CWBDataApp/report.html', {'dataAcceptedMessage':message['returnMessage'], 'allSubjects':allSubjects})
        else:
            return render(request, 'CWBDataApp/report.html', {'error_message':message['returnMessage'], 'allSubjects':allSubjects})


    return render(request, 'CWBDataApp/report.html', {'allSubjects':allSubjects})


###########################################################PRINT FORMS
def PrintForms(request):

    if request.method == 'POST':
        form = request.POST

        #Open PDF from pdf folder
        #webbrowser.open_new('file:///' + os.path.realpath('PdfForms/Mixing_Form.pdf'))
        return FileResponse(open( os.path.realpath('PdfForms/'+ form['filePrint'] +'.pdf'), 'rb'), content_type='application/pdf')
    return render(request, 'CWBDataApp/Forms.html')


###########################################################ADMIN
def admin(request):
    return reverse('admin:index')


#Add a given number of business days to a date, we ignore weekends
def addDate(startDate, numDays):
    currentDate = startDate

    while numDays > 0:
        currentDate += timedelta(days=1)
        weekday=currentDate.weekday()
        if weekday >= 5:
            continue
        else:
            numDays -= 1
    return currentDate

#Function to create and send email
def sendEmail(emailSubject, message):

    gmail_user = os.getenv('GMAIL_USER')
    gmail_password = os.getenv('GMAIL_PASSWORD')

    sent_from = gmail_user
    to = 'tom@greenwellplastics.ca'
    #to = 'ben10pickers@gmail.com'
    return_message = ''
    state = True

    msg = EmailMessage()
    msg.set_content(message)

    msg['Subject'] = emailSubject
    msg['From'] = sent_from
    msg['To'] = to

    try:
        context = ssl.create_default_context()
        smtp_server = smtplib.SMTP('smtp.gmail.com', 587)
        smtp_server.ehlo()
        smtp_server.starttls(context=context)
        smtp_server.ehlo()
        smtp_server.login(gmail_user, gmail_password)
        smtp_server.sendmail(sent_from, to, msg.as_string())
        smtp_server.close()
        returnMessage = "Email sent successfully! Thankyou for your input"
    except Exception as ex:
       state = False
       returnMessage = "Could not Send the email. Try again later"


    return {'state':state, 'returnMessage':returnMessage}

class signUpView(generic.CreateView):
    form_class = UserCreationForm
    success_url = reverse_lazy('CWBDataApp:login')
    template_name = 'registration/signup.html'
